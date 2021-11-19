from docx.enum.text import WD_BREAK
from docxcompose.composer import Composer
from docxtpl import DocxTemplate
import json


def report_handler(request):
    filepath = 'api/src/docs/'
    doc = DocxTemplate(filepath + "DeformationReportMain.docx")
    composer = Composer(doc)
    table_general = doc.tables[0]
    json_from_post = json.loads(request.POST["post"])
    list = ["Дорожная одежда", "Земляное полотно", "Мостовые сооружения"]
    types = [types["type"] for types in json_from_post["plots"]]

    object_table_to_add = DocxTemplate(filepath + "DeformationReportObjectTable.docx")
    for table in object_table_to_add.tables:
        is_table_suitable(table, list)
    for plot in json_from_post["plots"]:
        i = 1
        fill_general_table(json_from_post, plot, table_general, i)
        for table in object_table_to_add.tables:
            for row in table.rows[2:]:
                if row.cells[0].text.strip() == plot["type"]:
                    row.cells[2].text = plot["date"]
                    row.cells[4].text = plot["type"]
        i += 1

    clean_empty_rows(object_table_to_add)
    context = {"table_number": str(1)}
    object_table_to_add.render(context)
    composer.append(object_table_to_add)

    for par in doc.paragraphs:
        if "Таблица по объекту" in par.text:
            br = par.insert_paragraph_before("")
            run = br.add_run()
            run.add_break(WD_BREAK.PAGE)

    context = {'num': "7", "city": json_from_post['city'], "time": json_from_post['time'],
               "date": json_from_post['date'], "whos": json_from_post['whos']}
    doc.render(context)
    doc.save("gen.docx")


def fill_general_table(json_from_post, plot, table_general, i):
    row_cells = table_general.add_row().cells
    row_cells[0].text = str(i)
    row_cells[1].text = json_from_post["city"] + " " + plot["name"]
    row_cells[2].text = plot["date"]
    row_cells[3].text = plot["type"]


def is_table_suitable(table, list):
    for partition in list:
        if not table.rows[0].cells[0].text.strip().find(partition) == -1:
            return True
    table._element.getparent().remove(table._element)


def clean_empty_rows(object_table_to_add):
    for table in object_table_to_add.tables:
        for empty_row in table.rows:
            if empty_row.cells[2].text.strip() == "":  # and empty_row.cells[4].text.strip():
                empty_row._element.getparent().remove(empty_row._element)
