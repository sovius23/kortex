from docx.enum.text import WD_BREAK
from docxcompose.composer import Composer
from docxtpl import DocxTemplate
import json


def report_handler(request):
    table_dict = {
        "earth_canvas": {
            "Земляное полотно, полоса отвода": {"deformations": "Отдельные повреждения (деформации и разрушения)",
                                                "unsecured_drainage": "Необеспеченный водоотвод (застой воды)",
                                                "slopes": "Повреждения откосов насыпей и выемок",
                                                "damaged_drainage": "Повреждения системы водоотвода (водосбросы, "
                                                                    "дренажи, водоотводные канавы и др.)",
                                                "garbage": "Мусор и посторонние предметы",
                                                "deformation_of_colorization": "Дефекты элементов обозначения границ "
                                                                               "полосы отвода",
                                                "collapse_consequences": "Последствия обвалов, оползней, паводков, "
                                                                         "селевых потоков, пучин в результате "
                                                                         "несвоевременного проведения соответствующих "
                                                                         "мероприятий при содержании дороги", }
        },
        "road_clothes": {"Дорожная одежда": {"density_deformation": "Деформации и разрушения  с удалением материала",
                                             "colorization_deformation": "Деформации и разрушения без удаления материала",
                                             "subsidences": "Просадки",
                                             "potholes": "Выбоины",
                                             "coloring": "Выкрашивание",
                                             "peeling": "Шелушение",
                                             "breaks": "Проломы",
                                             "chips": "Сколы кромок",
                                             "huge_astringent_density": "Необработанные места выпотевания вяжущего",
                                             "profile_deformation": "Нарушение профиля, гребенка",
                                             "cracks": "Трещины",
                                             "deformated_seams": "Разрушенные и не заполненные мастикой "
                                                                 "деформационные швы на цементобетонном покрытии",
                                             "rut": "Колейность",
                                             "color_clothes_deformation": "Разрушение дорожной одежды на участках с "
                                                                          "пучинистыми и слабыми грунтами",
                                             "pollution_bands": "Полосы загрязнения у кромок покрытия",
                                             "other_objects": "Посторонние предметы на проезжей части",
                                             }
                         },
        "bridges": {"Мостовые сооружения (мостовое полотно)": {"pollution_of_bridges": "Загрязнение мостового полотна",
                                                               "stagnation_of_water": "Застой воды на проезжей части "
                                                                                      "и тротуарах",
                                                               "sidewalk_potholes": "Отдельные выбоины в покрытии "
                                                                                    "тротуаров, проломы в тротуарных "
                                                                                    "плитах",
                                                               "tubes_pollution": "Засорение водоотводных трубок и "
                                                                                  "окон в тротуарных блоках",
                                                               "": "",
                                                               }
                    },
        "road_fencing": {
            "Ограждения проезжей части": {"": "Повреждения отдельных секций металлического барьерного ограждения", }
        },
        "sidewalk_fenching": {"Перильные ограждения тротуаров": {"": "Повреждения отдельных секций перил", }
                              },
        "deformation_seams": {
            "Деформационные швы": {"": "Трещины в покрытии над деформационными швами, протечки в деформационных швах", }
        },
        "superstructures": {"Пролетные строения": {
            "supports_pollution": "Загрязнение насадок опор, опорных частей, лестничных сходов, перил и ограждений "
                                  "безопасности на мостовых сооружениях и подходах к ним",
            "garbage": "Мусор, загрязнение, растительность на пролетных строениях, конусах, под тротуарными блоками, "
                       "загрязнение подмостовой зоны",
            "bolts_defect": "Дефекты болтов и заклепок",
        }
        },
        "underbrigde_zone": {
            "Подмостовая зона": {"": "Нарушение поверхностей и структуры отдельных элементов конструкции", }
        },
        "tunnels": {"Тоннели, галереи, пешеходные переходы": {"lining": "Локальные повреждения обделки тоннеля",
                                                              "landslide": "Оползание грунта над порталами тоннеля",
                                                              "crosswalks": "Дефекты надземных (подземных) пешеходных "
                                                                            "переходов",
                                                              }
                    },
        "support_walls": {"Подпорные стенки": {"construction_defect": "Повреждение конструкции подпорных стенок",
                                               "washout": "Подмывы и размывы",
                                               }
                          },
        "cleaning_constructions": {"Очистные сооружения": {"garbage": "Мусор и посторонние предметы",
                                                           "water_treatment": "Нарушение системы водоочистки",
                                                           "slit": "Иловые отложения",
                                                           "flora": "Растительность",
                                                           "construction_defect": "Дефекты конструктивных элементов "
                                                                                  "очистных сооружений", }
                                   },
        "arrangement": {"Элементы обустройства автомобильных дорог": {
            "road_signs": "Дефекты дорожных знаков и табло с изменяющейся информацией. Дефекты табло с изменяющейся "
                          "информацией, затрудняющие ее восприятие",
            "guiding_device": "Дефекты направляющих устройств (дорожных сигнальных столбиков, дорожных тумб, буферов "
                              "и т.д.)",
            "road_fences": "Дефекты дорожных ограждений (в т.ч.пешеходных)",
            "traffic_lights": "Дефекты дорожных светофоров",
            "potholes": "Отдельные выбоины на покрытии тротуаров, пешеходных и велосипедных дорожек",
            "mirrors": "Дефекты дорожных зеркал",
            "curb": "Видимые повреждения бордюров",
            "stands_of_traffic_signs": "Дефекты стоек дорожных знаков (П, Г и Т-образные опоры)",
            "": "Дефекты остановочных пунктов общественного транспорта, площадок отдыха, площадок для остановки и "
                "кратковременной стоянки транспортных средств",
            "electricity_lines": "Дефекты линий наружного электроосвещения",
        }
        },
    }
    filepath = 'api/src/docs/DeformationReportMain.docx'
    main_doc = DocxTemplate(filepath)

    json_from_post = json.loads(request.POST["post"])
    plots = json_from_post["plots"][0]

    main_table(json_from_post, main_doc.tables[0])
    for table_name, defect_and_info in plots.items():
        if isinstance(defect_and_info, dict):
            for table in main_doc.tables[1:]:
                if table.rows[0].cells[0].text.strip() == json_to_rus(table_name, table_dict):
                    for rows in table.rows[2:]:
                        for defect, info in defect_and_info.items():
                            if rows.cells[0].text.strip() == json_to_rus(defect, table_dict):
                                rows.cells[2].text = info["date"]
                                rows.cells[3].text = str(info["number_of_annix"])
                                rows.cells[4].text = "Обнаружено"

    delete_empty_fields(main_doc.tables)
    context = {'num': "7", "city": json_from_post['city'], "time": json_from_post['time'],
               "date": json_from_post['date'], "whos": json_from_post['whos']}
    main_doc.render(context)
    add_page_break(main_doc)
    delete_empty_fields(main_doc.tables)
    return main_doc


def main_table(json_from_post, main_table):
    for plot in json_from_post["plots"]:
        row_cells = main_table.add_row().cells
        row_cells[0].text = str(plot["id"])
        row_cells[1].text = json_from_post["city"] + " " + plot["name"]
        row_cells[2].text = plot["date"]
        row_cells[3].text = plot["type"]


def delete_empty_fields(tables):
    for table in tables:
        for row in table.rows[2:]:
            if not row.cells[4].text:
                row._element.getparent().remove(row._element)
    for table in tables[1:]:
        if len(table.rows) == 2:
            table._element.getparent().remove(table._element)


def json_to_rus(json, table_dict):
    for name_table, name_rows in table_dict.items():
        if json == name_table:
            return list(table_dict[name_table].keys())[0]
        else:
            for name_row in name_rows.values():
                for json_name, rus_name in name_row.items():
                    if json == json_name:
                        return str(rus_name)


def add_page_break(main_doc):
    for par in main_doc.paragraphs:
        if "Таблица по объекту" in par.text:
            br = par.insert_paragraph_before(" ")
            run = br.add_run()
            run.add_break(WD_BREAK.PAGE)